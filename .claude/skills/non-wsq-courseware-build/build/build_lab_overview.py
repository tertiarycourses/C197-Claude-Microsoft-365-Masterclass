#!/usr/bin/env python3
"""Build labs/LAB-00-Overview.pdf — the one page that shows the whole HR story.

Learners and trainers open this first: what happens in each lab, which Office
app it uses, which file it opens, and what it hands to the next lab.
"""
import glob, importlib, os, re, shutil, subprocess, sys, tempfile

from docx import Document
from docx.shared import Pt, RGBColor

HERE = os.path.dirname(os.path.abspath(__file__)); sys.path.insert(0, HERE)
import course_data as C
import lab_files, prodoc

BRAND = RGBColor(0x1F, 0x6F, 0xEB)

def repo_root(start):
    d = start
    for _ in range(8):
        d = os.path.dirname(d)
        if os.path.isdir(os.path.join(d, "courseware")) and os.path.isdir(os.path.join(d, "labs")):
            return d
    raise RuntimeError("repo not found")

REPO = repo_root(HERE); LABS = os.path.join(REPO, "labs")
ASSETS = os.path.join(REPO, "courseware", "assets")

def acts():
    out = []
    for p in sorted(glob.glob(os.path.join(HERE, "data_domain[0-9]*.py"))):
        n = os.path.basename(p)[:-3]
        if re.fullmatch(r"data_domain\d+", n):
            out += getattr(importlib.import_module(n), "DOMAIN" + re.search(r"\d+", n).group())
    return sorted(out, key=lambda a: a["num"])

# The HR story: what the team is doing, and what each lab hands forward.
STORY = {
    0: ("Install the Claude add-in, Desktop, the Microsoft 365 connector and Claude for Chrome.", "A working setup and a record of what your account allows"),
    1: ("People are leaving. Shortlist replacements from 24 applicants.", "A shortlist of candidates"),
    2: ("Find out which team has the worst turnover, and decide what Claude may read.", "Warehouse is the problem team"),
    3: ("Learn how to ask Claude properly before touching real HR work.", "A request checklist you reuse"),
    4: ("Decide which roles to fill within the budget, and write the hiring plan.", "The FY2027 hiring plan"),
    5: ("Leavers cite inflexible hours. Draft the leave and flexible-work wording.", "Draft policy wording"),
    6: ("Work out what the hiring and the leavers actually cost.", "Headcount and cost analysis"),
    7: ("Present the quarter to the leadership team.", "A six-slide update"),
    8: ("Staff have questions about the new policy. Sort them and reply.", "A sorted inbox and one reply"),
    9: ("Pull the whole quarter into one summary for the Head of HR.", "A two-page people summary"),
    10: ("Automate the weekly update so next quarter runs itself.", "An automated weekly update"),
    11: ("Save the shortlisting method so the whole team works the same way.", "A reusable Skill"),
    12: ("Reach the team's real HR files in SharePoint, not just your own folder.", "A connected Microsoft 365 account"),
    13: ("Upload a shared standard so every deck looks the same.", "A shared slide-design skill"),
    14: ("Run a skill from Word that reads the Excel workbook for you.", "One skill working across two apps"),
}

def apps_for(a):
    folder = os.path.join(LABS, f"lab-{a['num']:02d}-{C.LAB_SLUGS[a['num']]}")
    seen = []
    for _label, app, _f, _t in lab_files.describe(folder):
        if app not in seen and app != "—":
            seen.append(app)
    return ", ".join(seen) or "—"

def main():
    doc = Document()
    normal = doc.styles["Normal"]; normal.font.name = "Arial"; normal.font.size = Pt(11)
    prodoc.style_headings(doc)
    prodoc.add_cover_page(doc, "LAB OVERVIEW", "One HR Quarter, Eleven Labs",
                          C.VERSION.lstrip("v"),
                          org_logo=os.path.join(ASSETS, "tertiary-infotech-logo.png"),
                          course_logo=None, course_code=C.COURSE_CODE)
    doc.add_heading("How the labs fit together", level=1)
    doc.add_paragraph(
        "Every lab is one step in a single HR quarter at Lumina Living. The finding in Lab 2 — that "
        "the warehouse team is losing people fastest — drives the hiring plan, the policy change, the "
        "cost analysis and the leadership update that follow.")
    doc.add_paragraph(
        "Each lab folder is still self-contained. Every file a lab needs is already inside it, so you "
        "can run any lab on its own, or join the course late, without being stuck.")

    t = doc.add_table(rows=1, cols=5); t.style = "Table Grid"
    heads = ("Lab", "What the HR team is doing", "App", "Main file you open", "What it gives the next lab")
    for cell, head in zip(t.rows[0].cells, heads):
        cell.text = ""; r = cell.paragraphs[0].add_run(head)
        r.bold = True; r.font.size = Pt(9)
        prodoc._shade_cell(cell, "EAF2FF")
    for a in acts():
        doing, gives = STORY.get(a["num"], ("", ""))
        folder = os.path.join(LABS, f"lab-{a['num']:02d}-{C.LAB_SLUGS[a['num']]}")
        main_file = "—"
        for _l, _app, fname, _tab in lab_files.describe(folder):
            if not fname.startswith("templates/"):
                main_file = fname; break
        cells = t.add_row().cells
        for cell, text in zip(cells, (str(a["num"]), doing, apps_for(a), main_file, gives)):
            cell.text = ""; run = cell.paragraphs[0].add_run(text); run.font.size = Pt(8.5)

    doc.add_paragraph()
    doc.add_heading("Which app each lab uses", level=2)
    doc.add_paragraph(
        "Excel — Labs 1, 2, 7, 9 and 11.  Word — Labs 3, 4, 5, 6 and 10.  PowerPoint — Lab 8.  "
        "Outlook is never opened: the staff messages in Lab 9 are supplied in a workbook so the lab "
        "runs on any computer.")
    prodoc.add_page_numbers(doc, left_text=f"{C.TITLE} · {C.COURSE_CODE} · Lab overview")

    tmpdocx = os.path.join(LABS, "COURSE-Lab-Overview.docx")
    doc.save(tmpdocx)
    with tempfile.TemporaryDirectory() as tmp:
        subprocess.run(["soffice", "--headless", "--convert-to", "pdf", "--outdir", tmp, tmpdocx],
                       check=True, stdout=subprocess.DEVNULL, stderr=subprocess.DEVNULL)
        made = glob.glob(os.path.join(tmp, "*.pdf"))
        dest = os.path.join(LABS, "COURSE-Lab-Overview.pdf")
        shutil.move(made[0], dest)
    os.remove(tmpdocx)
    print("Saved", dest)

if __name__ == "__main__":
    main()
