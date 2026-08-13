#!/usr/bin/env python3
"""Generate a standalone printable instruction PDF inside every lab folder.

Each lab folder gets LAB-NN-Instructions.pdf so a learner or trainer can open
one file and run the lab without the full Learner Guide.  Content comes from the
same canonical activity data that drives the LG, so the two cannot diverge.
"""

import glob
import importlib
import os
import re
import shutil
import subprocess
import sys
import tempfile

from docx import Document
from docx.shared import Pt, RGBColor
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

HERE = os.path.dirname(os.path.abspath(__file__))
sys.path.insert(0, HERE)
import course_data as C
import lab_files
import prodoc

BRAND = RGBColor(0x1F, 0x6F, 0xEB)
INKCODE = RGBColor(0x0B, 0x30, 0x60)


def repo_root(start):
    d = start
    for _ in range(8):
        d = os.path.dirname(d)
        if os.path.isdir(os.path.join(d, "courseware")) and os.path.isdir(os.path.join(d, "labs")):
            return d
    raise RuntimeError("Course repository not found")


def load_activities():
    acts = []
    for path in sorted(glob.glob(os.path.join(HERE, "data_domain[0-9]*.py"))):
        name = os.path.basename(path)[:-3]
        if not re.fullmatch(r"data_domain\d+", name):
            continue
        n = re.search(r"\d+", name).group()
        acts.extend(getattr(importlib.import_module(name), f"DOMAIN{n}"))
    return sorted(acts, key=lambda a: a["num"])


REPO = repo_root(HERE)
LABS = os.path.join(REPO, "labs")
ASSETS = os.path.join(REPO, "courseware", "assets")
ACTS = load_activities()
TOPICS = {t["num"]: t for t in C.TOPICS}

SHELL_HINTS = ("python", "pip", "cd ", "npm", "git ", "claude ", "./", "export ", "mkdir")


def looks_like_shell(payload):
    head = payload.strip().splitlines()[0].lower() if payload.strip() else ""
    return any(head.startswith(h) for h in SHELL_HINTS)


def code_para(doc, text, mono=True):
    for line in text.split("\n"):
        para = doc.add_paragraph()
        ppr = para._p.get_or_add_pPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:fill"), "EAF2FF")
        ppr.append(shd)
        para.paragraph_format.left_indent = Pt(10)
        para.paragraph_format.right_indent = Pt(10)
        # A blank line inside a prompt is a paragraph break, not a full empty
        # line — collapse it so the shaded block does not gape.
        blank = not line.strip()
        para.paragraph_format.space_before = Pt(0 if blank else 4)
        para.paragraph_format.space_after = Pt(0 if blank else 3)
        r = para.add_run(line)
        r.font.size = Pt(3 if blank else 9.5)
        r.font.color.rgb = INKCODE
        if mono:
            r.font.name = "Consolas"


def h3(doc, text):
    para = doc.add_paragraph()
    r = para.add_run(text)
    r.bold = True
    r.font.size = Pt(11)
    r.font.color.rgb = BRAND


def build_docx(a, out_docx):
    folder_name = f"lab-{a['num']:02d}-{C.LAB_SLUGS[a['num']]}"
    doc = Document()
    normal = doc.styles["Normal"]
    normal.font.name = "Arial"
    normal.font.size = Pt(11)
    prodoc.style_headings(doc)
    prodoc.add_cover_page(
        doc, f"LAB {a['num']:02d} INSTRUCTIONS", a["title"], C.VERSION.lstrip("v"),
        org_logo=os.path.join(ASSETS, "tertiary-infotech-logo.png"),
        course_logo=None, course_code=C.COURSE_CODE)

    doc.add_heading(f"Lab {a['num']} — {a['title']}", level=1)
    doc.add_paragraph(
        f"Topic {TOPICS[a['topic']]['code']} · {TOPICS[a['topic']]['title']} · "
        f"approximately {a.get('minutes', 20)} minutes")

    h3(doc, "Where the files for this lab are")
    doc.add_paragraph(
        "Every file this lab needs is inside one folder that came with your course "
        "materials. Open the labs folder, then open this folder inside it:")
    code_para(doc, f"labs/{folder_name}/", mono=True)
    doc.add_paragraph(
        "Do not look in any other lab folder. This lab is self-contained.")

    rows = lab_files.describe(os.path.join(LABS, folder_name))
    if rows:
        h3(doc, "Files you will use, and what the steps call them")
        doc.add_paragraph(
            "The steps below refer to these items in plain English. Use this table to find "
            "the exact file to open, and the exact sheet tab to click at the bottom of the "
            "Excel window.")
        table = doc.add_table(rows=1, cols=4)
        table.style = "Table Grid"
        for cell, head in zip(table.rows[0].cells, ("The steps call it", "App", "Open this file", "Then click")):
            cell.text = ""
            r = cell.paragraphs[0].add_run(head)
            r.bold = True
            r.font.size = Pt(9.5)
            prodoc._shade_cell(cell, "EAF2FF")
        for label, app, fname, tab in rows:
            cells = table.add_row().cells
            for cell, text in zip(cells, (label, app, fname, tab)):
                cell.text = ""
                run = cell.paragraphs[0].add_run(text)
                run.font.size = Pt(9.5)

    h3(doc, "Goal")
    doc.add_paragraph(a["desc"])

    h3(doc, "What you will produce")
    doc.add_paragraph(a["build"])

    h3(doc, "Before you start")
    for x in a["prerequisites"]:
        doc.add_paragraph(x, style="List Bullet")

    doc.add_heading("Steps", level=2)
    doc.add_paragraph(
        "Follow these in order. Where a shaded block appears, type or paste it exactly "
        "as shown into Claude, then read what Claude proposes before you accept anything.")
    for i, (instruction, payload) in enumerate(a["steps"], 1):
        para = doc.add_paragraph()
        para.paragraph_format.left_indent = Pt(18)
        para.paragraph_format.space_after = Pt(4)
        para.paragraph_format.keep_together = True
        para.paragraph_format.keep_with_next = bool(payload)
        r = para.add_run(f"{i}.  ")
        r.bold = True
        para.add_run(instruction)
        if payload:
            label = "Command to run:" if looks_like_shell(payload) else "Prompt to type into Claude:"
            lab = doc.add_paragraph()
            lab.paragraph_format.left_indent = Pt(18)
            lab.paragraph_format.space_after = Pt(2)
            lr = lab.add_run(label)
            lr.bold = True
            lr.font.size = Pt(10)
            code_para(doc, payload, mono=looks_like_shell(payload))

    doc.add_heading("Check your work", level=2)
    doc.add_paragraph(a["test"])

    doc.add_heading("If something goes wrong", level=2)
    for label, fix in a["troubleshooting"]:
        para = doc.add_paragraph(style="List Bullet")
        r = para.add_run(f"{label}. ")
        r.bold = True
        para.add_run(fix)

    doc.add_heading("Challenge", level=2)
    doc.add_paragraph(a["challenge"])
    doc.add_heading("Reflection", level=2)
    doc.add_paragraph(a["reflection"])

    prodoc.add_page_numbers(
        doc, left_text=f"{C.TITLE} · {C.COURSE_CODE} · Lab {a['num']:02d}")
    doc.save(out_docx)


def to_pdf(src_docx, out_dir):
    with tempfile.TemporaryDirectory() as tmp:
        subprocess.run(
            ["soffice", "--headless", "--convert-to", "pdf", "--outdir", tmp, src_docx],
            check=True, stdout=subprocess.DEVNULL, stderr=subprocess.DEVNULL)
        made = glob.glob(os.path.join(tmp, "*.pdf"))
        if not made:
            raise RuntimeError(f"PDF conversion produced nothing for {src_docx}")
        dest = os.path.join(out_dir, os.path.basename(made[0]))
        shutil.move(made[0], dest)
        return dest


def main():
    only = {int(x) for x in sys.argv[1:] if x.isdigit()}
    for a in ACTS:
        if only and a["num"] not in only:
            continue
        folder = os.path.join(LABS, f"lab-{a['num']:02d}-{C.LAB_SLUGS[a['num']]}")
        if not os.path.isdir(folder):
            print("Skipped (folder missing)", folder)
            continue
        # Remove any instruction PDF left behind by an earlier numbering.
        for old in glob.glob(os.path.join(folder, "LAB-*-Instructions.pdf")):
            if os.path.basename(old) != f"LAB-{a['num']:02d}-Instructions.pdf":
                os.remove(old)
                print("Removed stale", os.path.basename(old))
        stem = f"LAB-{a['num']:02d}-Instructions"
        docx_path = os.path.join(folder, stem + ".docx")
        build_docx(a, docx_path)
        pdf_path = to_pdf(docx_path, folder)
        # The DOCX is only an intermediate; learners open the PDF.
        os.remove(docx_path)
        print("Saved", pdf_path)


if __name__ == "__main__":
    main()
