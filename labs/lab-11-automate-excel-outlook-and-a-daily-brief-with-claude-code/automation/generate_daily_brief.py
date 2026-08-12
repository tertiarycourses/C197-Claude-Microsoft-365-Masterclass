#!/usr/bin/env python3
"""Generate a source-linked fictional daily management brief."""
import argparse, json
from docx import Document
from openpyxl import load_workbook

p=argparse.ArgumentParser(); p.add_argument("--workbook",required=True); p.add_argument("--mail",required=True); p.add_argument("--template",required=True); p.add_argument("--output",required=True); args=p.parse_args()
wb=load_workbook(args.workbook,data_only=False); ws=wb["Management_Control"]
exceptions=[]
for r in range(2,ws.max_row+1):
    status=ws.cell(r,8).value
    if status and "Attention" in str(status): exceptions.append((ws.cell(r,1).value,ws.cell(r,2).value,f"Management_Control!H{r}"))
mail=json.load(open(args.mail,encoding="utf-8")); doc=Document(args.template)
doc.add_heading("Today's KPI exceptions",level=1)
for month,metric,cite in exceptions[:8]: doc.add_paragraph(f"{month} · {metric} — review required ({cite})",style="List Bullet")
doc.add_heading("Planning decisions and actions",level=1)
for item in mail: doc.add_paragraph(f"{item['summary']} ({item['citation']})",style="List Bullet")
doc.add_heading("Reviewer checklist",level=1)
for item in ["Workbook cells checked","Message citations opened","Recipients and attachments verified","Human approval recorded"]: doc.add_paragraph(item,style="List Bullet")
doc.save(args.output); print(args.output)
